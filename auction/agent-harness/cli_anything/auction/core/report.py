"""Report generation for auction cases.

Generates DOCX and text reports from analysis results.
"""

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional


def generate_text_report(analysis: dict) -> str:
    """Generate a plain text report from analysis results."""
    rec = analysis.get("recommendation", "?")
    lines = [
        f"{'='*60}",
        f"AUCTION ANALYSIS REPORT",
        f"{'='*60}",
        f"Case:           {analysis.get('case_number', '?')}",
        f"Address:        {analysis.get('address', '?')}",
        f"Plaintiff:      {analysis.get('plaintiff', '?')}",
        f"{'─'*60}",
        f"Judgment:        ${analysis.get('judgment_amount', 0):,.2f}",
        f"ARV:             ${analysis.get('arv', 0):,.2f}",
        f"Repairs:         ${analysis.get('repairs', 0):,.2f}",
        f"Max Bid:         ${analysis.get('max_bid', 0):,.2f}",
        f"Bid/Judgment:    {analysis.get('bid_ratio', 0):.1%}",
        f"{'─'*60}",
        f"RECOMMENDATION:  {rec}",
        f"{'='*60}",
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
    ]
    return "\n".join(lines)


def generate_report(analysis: dict, fmt: str = "text", output_path: Optional[str] = None) -> dict:
    """Generate a report in the specified format.

    Formats: text, json, docx
    """
    if fmt == "json":
        content = json.dumps(analysis, indent=2, default=str)
    elif fmt == "docx":
        content = _generate_docx(analysis, output_path)
        if isinstance(content, dict):
            return content  # Already a result dict
    else:
        content = generate_text_report(analysis)

    if output_path:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content)
        return {
            "format": fmt,
            "path": str(path),
            "case_number": analysis.get("case_number"),
            "recommendation": analysis.get("recommendation"),
            "size_bytes": path.stat().st_size,
        }

    return {
        "format": fmt,
        "case_number": analysis.get("case_number"),
        "recommendation": analysis.get("recommendation"),
        "content": content,
    }


def _generate_docx(analysis: dict, output_path: Optional[str]) -> dict:
    """Generate DOCX report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading("Auction Analysis Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case info
        doc.add_heading("Case Details", level=2)
        doc.add_paragraph(f"Case Number: {analysis.get('case_number', '?')}")
        doc.add_paragraph(f"Address: {analysis.get('address', '?')}")
        doc.add_paragraph(f"Plaintiff: {analysis.get('plaintiff', '?')}")

        # Financial analysis
        doc.add_heading("Financial Analysis", level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = "Light List"
        rows = table.rows
        data = [
            ("Judgment Amount", f"${analysis.get('judgment_amount', 0):,.2f}"),
            ("After-Repair Value (ARV)", f"${analysis.get('arv', 0):,.2f}"),
            ("Estimated Repairs", f"${analysis.get('repairs', 0):,.2f}"),
            ("Maximum Bid", f"${analysis.get('max_bid', 0):,.2f}"),
            ("Bid/Judgment Ratio", f"{analysis.get('bid_ratio', 0):.1%}"),
            ("RECOMMENDATION", analysis.get("recommendation", "?")),
        ]
        for i, (label, value) in enumerate(data):
            rows[i].cells[0].text = label
            rows[i].cells[1].text = value

        # Save
        if output_path:
            path = Path(output_path)
            path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
            return {
                "format": "docx",
                "path": str(path),
                "case_number": analysis.get("case_number"),
                "recommendation": analysis.get("recommendation"),
                "size_bytes": path.stat().st_size,
            }
        return {"format": "docx", "error": "No output path specified"}

    except ImportError:
        # Fallback to text if python-docx not installed
        return generate_report(analysis, fmt="text", output_path=output_path)


def batch_reports(analyses: list[dict], output_dir: str, fmt: str = "text") -> dict:
    """Generate reports for multiple cases."""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    generated = []
    for analysis in analyses:
        case = analysis.get("case_number", "unknown").replace("-", "_")
        ext = "docx" if fmt == "docx" else "txt"
        path = str(Path(output_dir) / f"report_{case}.{ext}")
        result = generate_report(analysis, fmt=fmt, output_path=path)
        generated.append(result)

    return {
        "output_dir": output_dir,
        "format": fmt,
        "total": len(analyses),
        "generated": len(generated),
        "reports": generated,
    }
